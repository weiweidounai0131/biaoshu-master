#!/usr/bin/env python3
"""主标DOCX的轻量结构与残留词审计；不替代人工语义和WPS视觉校验。"""

import argparse
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document


RESIDUE_TERMS = [
    "内部会议", "根据会议", "会议讨论", "评分表要求", "为了得分", "需求书提到",
    "本批次", "上半部分", "下半部分", "后续章节", "后续部分", "本章（续）",
    "本部分在", "黄色标注", "投标供应商", "编写说明", "页数要求",
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--max-heading-depth", type=int, default=6)
    args = parser.parse_args()

    doc = Document(args.docx)
    headings = []
    residues = []
    body = Counter()
    images = []

    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        match = re.fullmatch(r"Heading (\d+)", style)
        if match:
            headings.append({"paragraph": index, "level": int(match.group(1)), "text": text})
        elif text and not text.startswith(("【图片占位：", "【图片规划占位：")):
            body[text] += 1
        if text.startswith(("【图片占位：", "【图片规划占位：")):
            images.append({"paragraph": index, "text": text})
        for term in RESIDUE_TERMS:
            if term in text:
                residues.append({"paragraph": index, "term": term, "text": text})

    duplicate_body = [
        {"count": count, "text": text}
        for text, count in body.items()
        if count > 1 and len(text) >= 24
    ]
    depth_over = [item for item in headings if item["level"] > args.max_heading_depth]
    h1 = [item["text"] for item in headings if item["level"] == 1]

    result = {
        "file": str(args.docx),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "heading_levels": dict(sorted(Counter(item["level"] for item in headings).items())),
        "h1": h1,
        "heading_depth_over_limit": depth_over,
        "image_placeholders": images,
        "residue_hits": residues,
        "duplicate_body": duplicate_body,
        "finding_count": len(depth_over) + len(residues) + len(duplicate_body),
    }
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(payload, encoding="utf-8")
    print(payload)


if __name__ == "__main__":
    main()
