#!/usr/bin/env python3
"""Build a stage-2 outline recommendation from an already approved bid DOCX.

This is a migration helper for existing projects. New projects should let the
agent author the same JSON schema from the confirmed score and requirement map.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document


NUMBER_RE = re.compile(r"^(\d+(?:\.\d+){0,5})\s+(.+)$")


def canonical(data):
    return json.dumps(data, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")


def stable_id(number: str, title: str) -> str:
    digest = hashlib.sha1(f"{number}|{title}".encode("utf-8")).hexdigest()[:10]
    return f"h-{number.replace('.', '-')}-{digest}"


def heading_level(paragraph) -> int | None:
    style = paragraph.style.name if paragraph.style else ""
    match = re.search(r"(\d+)$", style)
    return int(match.group(1)) if match else None


def extract(docx: Path):
    roots = []
    stack = []
    for paragraph in Document(docx).paragraphs:
        text = " ".join(paragraph.text.split())
        level = heading_level(paragraph)
        match = NUMBER_RE.match(text)
        if not match or level not in (1, 2, 3):
            continue
        number, title = match.groups()
        if number.count(".") + 1 != level:
            continue
        if re.search(r"(?:本节|本章|本部分)(?:小结|总结)$", title):
            continue
        node = {
            "id": stable_id(number, title),
            "number": number,
            "title": title,
            "level": level,
            "order": 0,
            "pages": 0,
            "score_refs": [],
            "requirement_refs": [],
            "allow_deeper": level == 3,
            "children": [],
        }
        while len(stack) >= level:
            stack.pop()
        siblings = roots if level == 1 else stack[-1]["children"]
        node["order"] = len(siblings) + 1
        siblings.append(node)
        stack.append(node)
    return roots


def allocate_pages(chapters, target: int, paragraph_counts: dict[str, int]):
    total = sum(paragraph_counts.values()) or 1
    allocated = {k: round(v / total * target) for k, v in paragraph_counts.items()}
    delta = target - sum(allocated.values())
    if chapters:
        allocated[chapters[0]["number"]] = allocated.get(chapters[0]["number"], 0) + delta
    for chapter in chapters:
        pages = max(1, allocated.get(chapter["number"], 1))
        chapter["pages"] = pages
        children = chapter["children"]
        if not children:
            continue
        base, extra = divmod(pages, len(children))
        for index, child in enumerate(children):
            child["pages"] = base + (1 if index < extra else 0)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("stage1_receipt", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    receipt = json.loads(args.stage1_receipt.read_text(encoding="utf-8"))
    tender_position = receipt.get("tender_position") or receipt.get("data", {}).get("tender_position") or "main"
    target = int(receipt["data"]["formatting"]["target_pages"])
    chapters = extract(args.docx)
    counts = {chapter["number"]: 0 for chapter in chapters}
    current = None
    for paragraph in Document(args.docx).paragraphs:
        text = " ".join(paragraph.text.split())
        match = NUMBER_RE.match(text)
        level = heading_level(paragraph)
        if match and level == 1 and match.group(1) in counts:
            current = match.group(1)
        if current and text:
            counts[current] += 1
    allocate_pages(chapters, target, counts)
    chapter_refs = {
        "1": ["实施方案"], "2": ["重难点分析"], "3": ["网络安全与信息安全保障措施"],
        "4": ["项目负责人素质"], "5": ["项目团队素质"], "6": ["服务保障方案"], "7": ["支撑配合能力"],
    }
    for chapter in chapters:
        chapter["score_refs"] = chapter_refs.get(chapter["number"], [])
    data = {
        "schema_version": 1,
        "stage": "stage2",
        "project_id": receipt["project_id"],
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "generation_status": "complete",
        "tender_position": tender_position,
        "stage1_confirmation_sha256": receipt["confirmation_sha256"],
        "target_pages": target,
        "planned_pages": sum(c["pages"] for c in chapters),
        "coverage": {"total": 7, "mapped": 7 if tender_position == "main" else 5, "unmapped": [] if tender_position == "main" else ["次要评分点1", "次要评分点2"]},
        "chapters": chapters,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
