#!/usr/bin/env python3
"""Export one approved structured bid source to a formatted DOCX.

This is a deterministic local renderer. It never invents source text, sends
content to a model, generates images, or inserts image binaries. Any image
block remains a clearly labelled planning placeholder in the Word file.
"""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from . import protocol
except ImportError:
    import protocol


FONT_NAME = "仿宋_GB2312"
FONT_SIZE = Pt(14)
RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run: Any, *, bold: bool = False, color: RGBColor = BLACK) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = FONT_SIZE
    run.bold = bold
    run.font.color.rgb = color


def set_paragraph_format(paragraph: Any, *, heading: bool = False, list_item: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(10) if heading else Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.25
    fmt.first_line_indent = Pt(0) if heading or list_item else Pt(28)
    if heading:
        fmt.keep_with_next = True


def add_text_paragraph(document: Document, text: str, *, heading_level: int | None = None, list_item: bool = False, color: RGBColor = BLACK) -> Any:
    heading = heading_level is not None
    paragraph = document.add_paragraph(style=f"Heading {heading_level}" if heading else None)
    set_paragraph_format(paragraph, heading=heading, list_item=list_item)
    run = paragraph.add_run(text)
    set_run_font(run, bold=heading, color=color)
    return paragraph


def configure_heading_styles(document: Document) -> None:
    """Keep real Word heading semantics while applying the bid's typography."""
    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = FONT_SIZE
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True


def set_cell_border(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")


def add_table(document: Document, columns: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    data = [columns, *rows]
    for row_index, row_values in enumerate(data):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(paragraph)
            run = paragraph.add_run(value)
            set_run_font(run, bold=(row_index == 0))


def add_image_placeholder(document: Document, block: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(paragraph)
    run = paragraph.add_run(f"【图片占位：{block['figure_no']} {block['name']}】\n{block['note']}")
    set_run_font(run, bold=True, color=RED)


def add_material_placeholder(document: Document, block: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(f"【待补材料：{block['label']}】{block['note']}")
    set_run_font(run, bold=True)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def document_from_source(source: dict[str, Any]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = FONT_SIZE
    normal.font.color.rgb = BLACK
    configure_heading_styles(document)
    if "Bid Body" not in document.styles:
        style = document.styles.add_style("Bid Body", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = FONT_SIZE
        style.font.color.rgb = BLACK

    chapter_by_id = {chapter["id"]: chapter for chapter in source["chapters"]}
    seen_chapters: set[str] = set()
    for block in source["blocks"]:
        chapter_id = block["chapter_id"]
        if chapter_id not in seen_chapters:
            chapter = chapter_by_id[chapter_id]
            level = int(chapter.get("level", 1))
            if not 1 <= level <= 6:
                raise ValueError("Word导出章节标题层级必须为1至6")
            add_text_paragraph(document, f"{chapter['number']} {chapter['title']}", heading_level=level)
            seen_chapters.add(chapter_id)
        kind = block["type"]
        if kind == "heading":
            level = int(block["level"])
            if not 2 <= level <= 6:
                raise ValueError("Word导出小标题层级必须为2至6")
            add_text_paragraph(document, f"{block['number']} {block['title']}", heading_level=level)
        elif kind == "paragraph":
            add_text_paragraph(document, block["text"])
        elif kind == "list":
            for item in block["items"]:
                add_text_paragraph(document, f"· {item}", list_item=True)
        elif kind == "table":
            add_table(document, block["columns"], block["rows"])
        elif kind == "image_placeholder":
            add_image_placeholder(document, block)
        elif kind == "material_placeholder":
            add_material_placeholder(document, block)
        elif kind == "page_break":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    return document


def validate_docx(path: Path, source: dict[str, Any], calibration_ratio: float = 1.0) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("Word导出文件结构无效")
    reopened = Document(path)
    expected_headings = {f"{item['number']} {item['title']}": int(item.get("level", 1)) for item in source["chapters"]}
    expected_headings.update({f"{item['number']} {item['title']}": int(item["level"]) for item in source["blocks"] if item["type"] == "heading"})
    heading_count = sum(
        1 for paragraph in reopened.paragraphs
        if paragraph.text.strip() in expected_headings
        and paragraph.style.name == f"Heading {expected_headings[paragraph.text.strip()]}"
    )
    if heading_count != len(expected_headings):
        raise ValueError("Word导出文件缺少章节标题")
    estimate = protocol.estimate_source_pages(source, calibration_ratio)
    structure = protocol.source_structure_metrics(source)
    lower_bound, upper_bound = protocol.page_bounds(source["planned_pages"])
    return {
        "schema_version": 1,
        "kind": "bid_delivery_word_validation",
        "checked_at": protocol.utc_now(),
        "checks": {
            "docx_package_valid": True,
            "docx_reopenable": True,
            "chapter_heading_count": heading_count,
            "table_count": len(reopened.tables),
            "image_placeholder_count": sum(1 for block in source["blocks"] if block["type"] == "image_placeholder"),
            "material_placeholder_count": sum(1 for block in source["blocks"] if block["type"] == "material_placeholder"),
            "structure": structure,
        },
        "page_verification": {
            "status": "pending_wps_check",
            "planned_pages": source["planned_pages"],
            # WPS pagination belongs to this exported artifact, never to the
            # source draft. A revised export must not inherit an old mismatch.
            "actual_pages": None,
            "estimated_pages": estimate["estimated_pages"],
            "raw_estimated_pages": estimate["raw_estimated_pages"],
            "calibration_ratio": estimate["calibration_ratio"],
            "estimated_units": estimate["estimated_units"],
            "chapter_estimates": estimate["chapter_estimates"],
            "allowed_min_pages": lower_bound,
            "allowed_max_pages": upper_bound,
            "note": "预计页数仅用于发现内容不足；最终须以WPS实际页数复核，并处于允许范围内。",
        },
    }


def export_word(project_dir: Path, batch_id: str) -> dict[str, Any]:
    manifest, batch, source = protocol.load_batch_source_for_export(project_dir, batch_id)
    output_path = protocol.delivery_dir(project_dir) / batch["export_path"]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document_from_source(source).save(output_path)
    calibration = protocol.load_page_calibration(project_dir)
    validation = validate_docx(output_path, source, calibration["ratio"])
    validation.update({
        "project_id": manifest["project_id"], "batch_id": batch_id,
        "source_sha256": protocol.sha256_file(protocol.delivery_dir(project_dir) / batch["source_path"]),
        "export_sha256": protocol.sha256_file(output_path),
    })
    protocol.atomic_write_json(protocol.delivery_dir(project_dir) / protocol.RESULTS_DIR_NAME / f"word-batch-{batch['order']:02d}-validation.json", validation)
    updated = protocol.register_batch_artifacts(project_dir, batch_id)
    return {"manifest": updated, "output": str(output_path), "validation": validation}


def main() -> int:
    parser = argparse.ArgumentParser(description="导出一批主标Word")
    parser.add_argument("project_dir", type=Path)
    parser.add_argument("--batch", required=True, dest="batch_id")
    args = parser.parse_args()
    result = export_word(args.project_dir.expanduser().resolve(), args.batch_id)
    print(json.dumps({"output": result["output"], "status": result["manifest"]["status"], "validation": result["validation"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
